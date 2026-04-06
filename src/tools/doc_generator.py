"""Document generator tool for creating DOCX and PDF files."""
from __future__ import annotations

import logging
from enum import Enum
from pathlib import Path

logger = logging.getLogger(__name__)


class DocFormat(str, Enum):
    """Supported document output formats."""

    DOCX = "docx"
    PDF = "pdf"
    MARKDOWN = "markdown"


class DocGenerator:
    """Generate documents in DOCX or PDF format from markdown content.

    Uses python-docx for DOCX generation. PDF generation requires
    an external tool (pandoc or weasyprint) and is a stub for now.
    """

    def __init__(self, output_dir: str | Path) -> None:
        """Initialize the document generator.

        Args:
            output_dir: Directory where generated documents are saved.
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def create_docx(
        self,
        content: str,
        filename: str,
        title: str | None = None,
    ) -> Path:
        """Create a DOCX file from markdown content.

        Args:
            content: Markdown text content.
            filename: Output filename (without extension).
            title: Optional document title.

        Returns:
            Path to the generated DOCX file.

        Raises:
            ImportError: If python-docx is not installed.
        """
        try:
            from docx import Document  # type: ignore[import-untyped]
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DOCX generation. "
                "Install with: pip install python-docx"
            ) from exc

        doc = Document()
        if title:
            doc.add_heading(title, level=0)
        paragraphs = content.split("\n\n")
        for para in paragraphs:
            text = para.strip()
            if not text:
                continue
            if text.startswith("# "):
                doc.add_heading(text[2:], level=1)
            elif text.startswith("## "):
                doc.add_heading(text[3:], level=2)
            elif text.startswith("### "):
                doc.add_heading(text[4:], level=3)
            else:
                doc.add_paragraph(text)
        out_path = self.output_dir / f"{filename}.docx"
        doc.save(str(out_path))
        logger.info("Created DOCX: %s", out_path)
        return out_path

    def create_pdf(
        self,
        content: str,
        filename: str,
        title: str | None = None,
    ) -> Path:
        """Create a PDF file from markdown content.

        This is a stub that writes the markdown to a .md file.
        Full PDF generation requires pandoc or weasyprint.

        Args:
            content: Markdown text content.
            filename: Output filename (without extension).
            title: Optional document title.

        Returns:
            Path to the generated file (markdown fallback).
        """
        logger.warning(
            "PDF generation is a stub. "
            "Writing markdown file instead. "
            "Install pandoc for full PDF support."
        )
        out_path = self.output_dir / f"{filename}.md"
        full_content = content
        if title:
            full_content = f"# {title}\n\n{content}"
        out_path.write_text(full_content, encoding="utf-8")
        return out_path

    def create(
        self,
        content: str,
        filename: str,
        fmt: DocFormat,
        title: str | None = None,
    ) -> Path:
        """Create a document in the specified format.

        Args:
            content: Markdown text content.
            filename: Output filename (without extension).
            fmt: Desired output format.
            title: Optional document title.

        Returns:
            Path to the generated file.

        Raises:
            ValueError: If format is not supported.
        """
        if fmt == DocFormat.DOCX:
            return self.create_docx(content, filename, title)
        if fmt == DocFormat.PDF:
            return self.create_pdf(content, filename, title)
        if fmt == DocFormat.MARKDOWN:
            out_path = self.output_dir / f"{filename}.md"
            out_path.write_text(content, encoding="utf-8")
            return out_path
        raise ValueError(f"Unsupported format: {fmt}")
